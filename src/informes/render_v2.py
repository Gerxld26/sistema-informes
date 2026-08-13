from pathlib import Path

from docx.shared import Pt
from docxtpl import DocxTemplate

from . import config, textos

RUTA_PLANTILLA_EQUIPO = config.RAIZ / "plantillas" / "plantilla_informe_tecnico_equipo.docx"
RUTA_PLANTILLA_CONSUMIBLE = config.RAIZ / "plantillas" / "plantilla_informe_tecnico_consumible.docx"
RUTA_PLANTILLA_PERIFERICO = config.RAIZ / "plantillas" / "plantilla_informe_tecnico_periferico.docx"

RUTAS_PLANTILLA = {
    config.ESTILO_EQUIPO: RUTA_PLANTILLA_EQUIPO,
    config.ESTILO_CONSUMIBLE: RUTA_PLANTILLA_CONSUMIBLE,
    config.ESTILO_PERIFERICO: RUTA_PLANTILLA_PERIFERICO,
}


def _construir_tabla(tpl: DocxTemplate, columnas: list[str], filas: list[dict]):
    subdoc = tpl.new_subdoc()
    tabla = subdoc.add_table(rows=1, cols=len(columnas))
    tabla.style = "Table Grid"
    tabla.autofit = True

    encabezado = tabla.rows[0].cells
    for i, col in enumerate(columnas):
        encabezado[i].text = config.ETIQUETAS_COLUMNA.get(col, col.upper())
        encabezado[i].paragraphs[0].runs[0].bold = True
        encabezado[i].paragraphs[0].runs[0].font.size = Pt(9)

    for fila_datos in filas:
        celdas = tabla.add_row().cells
        for i, col in enumerate(columnas):
            celdas[i].text = str(fila_datos.get(col, "") or "")
            for run in celdas[i].paragraphs[0].runs:
                run.font.size = Pt(9)

    return subdoc


def generar_informe_tecnico_v2(tramite: dict, equipos: list[dict], carpeta_salida: Path = config.RUTA_SALIDA) -> Path:
    tipo_bien = tramite["tipo_bien"]
    info_bien = config.TIPOS_BIEN.get(tipo_bien, config.TIPOS_BIEN["CPU"])
    estilo = info_bien["estilo"]
    columnas = info_bien["columnas"]

    ruta_plantilla = RUTAS_PLANTILLA[estilo]
    tpl = DocxTemplate(str(ruta_plantilla))

    info_tipo = config.TEXTOS_TIPO_TRAMITE[tramite["tipo_tramite"]]

    contexto = {
        "numero_informe": str(int(tramite["numero_informe"])).zfill(3),
        "anio": int(tramite["anio"]),
        "iniciales_tecnicos": textos.iniciales(tramite["tecnico_nombre"]),
        "siglas_sede": tramite["siglas_sede"],
        "destinatario_nombre": tramite["destinatario_nombre"],
        "destinatario_cargo": tramite["destinatario_cargo"],
        "tecnico_nombre": tramite["tecnico_nombre"],
        "tecnico_cargo": tramite["tecnico_cargo"],
        "asunto_linea1": tramite["asunto_linea1"],
        "asunto_linea2": tramite.get("asunto_linea2", ""),
        "fecha": tramite["fecha"],
        "intro": tramite["intro"],
        "area": tramite["area"],
        "usuario": tramite["usuario"],
        "etiqueta_responsable": tramite.get("etiqueta_responsable", "JEFATURA"),
        "nombre_responsable": tramite.get("nombre_responsable", ""),
        "jefe_ti_nombre": tramite["jefe_ti_nombre"],
        "jefe_ti_cargo": tramite["jefe_ti_cargo"],
    }

    if estilo == config.ESTILO_PERIFERICO:
        contexto.update({
            "grupos_problema": tramite.get("grupos_problema", []),
            "conclusion": tramite.get("conclusion", ""),
            "recomendaciones": tramite.get("recomendaciones", []),
        })
    else:
        contexto.update({
            "tiene_descripcion_problema": info_tipo["tiene_descripcion_problema"],
            "descripcion_problema": tramite.get("descripcion_problema", []),
            "tabla_equipos": _construir_tabla(tpl, columnas, equipos),
            "texto_recomendacion": tramite.get("texto_recomendacion", ""),
            "especificaciones": tramite.get("especificaciones", []),
        })

    tpl.render(contexto)

    id_tramite = textos.quitar_tildes(tramite["id_tramite"])
    nombre_archivo = (
        f"Informe_Tecnico_{contexto['numero_informe']}-{contexto['anio']}_"
        f"{contexto['iniciales_tecnicos']}_{id_tramite}.docx"
    )
    carpeta_salida = Path(carpeta_salida)
    carpeta_salida.mkdir(parents=True, exist_ok=True)
    ruta_out = carpeta_salida / nombre_archivo
    tpl.save(str(ruta_out))
    return ruta_out
