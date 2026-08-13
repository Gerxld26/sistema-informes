from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RAIZ = Path(__file__).resolve().parent.parent
RUTA_SALIDA = RAIZ / "plantillas" / "plantilla_informe_tecnico_equipo.docx"

doc = Document()

section = doc.sections[0]
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)


def sin_bordes(tabla):
    tblPr = tabla._tbl.tblPr
    bordes = OxmlElement("w:tblBorders")
    for lado in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{lado}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        bordes.append(el)
    tblPr.append(bordes)


def celda(cell, texto, negrita_lineas=None, tam=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    lineas = texto.split("\n")
    for idx, linea in enumerate(lineas):
        if idx > 0:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
        r = p.add_run(linea)
        r.font.size = Pt(tam)
        if negrita_lineas and idx in negrita_lineas:
            r.bold = True


def parrafo(texto="", espacio_despues=6, sangria=None, negrita=False, es_tag=False, vineta=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0 if es_tag else espacio_despues)
    if sangria:
        p.paragraph_format.left_indent = Cm(sangria)
    if texto:
        prefijo = f"{vineta}  " if vineta else ""
        r = p.add_run(prefijo + texto)
        r.bold = negrita
    return p


titulo = parrafo(espacio_despues=10)
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = titulo.add_run(
    "INFORME TECNICO N.\u00b0 {{ numero_informe }}-{{ anio }} {{ iniciales_tecnicos }}/{{ siglas_sede }}"
)
run.bold = True
run.underline = True
run.font.size = Pt(11)

tabla_header = doc.add_table(rows=4, cols=3)
tabla_header.autofit = False
anchos = [Cm(2.5), Cm(0.5), Cm(11.5)]
for fila in tabla_header.rows:
    for i, c in enumerate(fila.cells):
        c.width = anchos[i]

celda(tabla_header.rows[0].cells[0], "A")
celda(tabla_header.rows[0].cells[1], ":")
celda(tabla_header.rows[0].cells[2], "{{ destinatario_nombre }}\n{{ destinatario_cargo }}", {1})

celda(tabla_header.rows[1].cells[0], "DE")
celda(tabla_header.rows[1].cells[1], ":")
celda(tabla_header.rows[1].cells[2], "{{ tecnico_nombre }}\n{{ tecnico_cargo }}", {1})

celda(tabla_header.rows[2].cells[0], "ASUNTO")
celda(tabla_header.rows[2].cells[1], ":")
celda(tabla_header.rows[2].cells[2], "{{ asunto_linea1 }}\n{{ asunto_linea2 }}", {0, 1})

celda(tabla_header.rows[3].cells[0], "FECHA")
celda(tabla_header.rows[3].cells[1], ":")
celda(tabla_header.rows[3].cells[2], "{{ fecha }}")

sin_bordes(tabla_header)

parrafo(espacio_despues=6)

linea = parrafo(espacio_despues=8)
pPr = linea._p.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
inferior = OxmlElement("w:bottom")
inferior.set(qn("w:val"), "single")
inferior.set(qn("w:sz"), "6")
inferior.set(qn("w:space"), "1")
inferior.set(qn("w:color"), "auto")
pBdr.append(inferior)
pPr.append(pBdr)

parrafo("{{ intro }}", espacio_despues=8)

parrafo("{% if tiene_descripcion_problema %}", es_tag=True)
parrafo("DESCRIPCION DEL PROBLEMA:", espacio_despues=4, negrita=True)
parrafo("{% for punto in descripcion_problema %}", es_tag=True)
parrafo("{{ punto }}", espacio_despues=2, sangria=1, vineta="\u2022")
parrafo("{% endfor %}", es_tag=True)
parrafo(espacio_despues=6)
parrafo("{% endif %}", es_tag=True)

parrafo("{{p tabla_equipos }}", es_tag=True)

parrafo(espacio_despues=6)
parrafo("RECOMENDACIONES:", espacio_despues=4, negrita=True)
parrafo("{{ texto_recomendacion }}", espacio_despues=4)
parrafo("{% for spec in especificaciones %}", es_tag=True)
parrafo("{{ spec }}", espacio_despues=2, sangria=1, vineta="\u2022")
parrafo("{% endfor %}", es_tag=True)

parrafo(espacio_despues=6)

tabla_pie = doc.add_table(rows=3, cols=3)
w2 = [Cm(2.3), Cm(0.5), Cm(11.7)]
for fila in tabla_pie.rows:
    for i, c in enumerate(fila.cells):
        c.width = w2[i]
celda(tabla_pie.rows[0].cells[0], "\u00c1REA")
celda(tabla_pie.rows[0].cells[1], ":")
celda(tabla_pie.rows[0].cells[2], "{{ area }}")
celda(tabla_pie.rows[1].cells[0], "USUARIO")
celda(tabla_pie.rows[1].cells[1], ":")
celda(tabla_pie.rows[1].cells[2], "{{ usuario }}")
celda(tabla_pie.rows[2].cells[0], "{{ etiqueta_responsable }}")
celda(tabla_pie.rows[2].cells[1], ":")
celda(tabla_pie.rows[2].cells[2], "{{ nombre_responsable }}")
sin_bordes(tabla_pie)

parrafo(espacio_despues=6)

firma1 = parrafo("_______________________", espacio_despues=2)
firma1.alignment = WD_ALIGN_PARAGRAPH.CENTER

firma2 = parrafo(espacio_despues=2)
firma2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = firma2.add_run("{{ jefe_ti_nombre }}")
r.bold = True

firma3 = parrafo()
firma3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = firma3.add_run("{{ jefe_ti_cargo }}")
r3.bold = True

RUTA_SALIDA.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(RUTA_SALIDA))
print(f"Plantilla creada en {RUTA_SALIDA}")
