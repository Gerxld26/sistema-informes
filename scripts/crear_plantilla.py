from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RAIZ = Path(__file__).resolve().parent.parent
RUTA_SALIDA = RAIZ / "plantillas" / "plantilla_informe_tecnico.docx"

doc = Document()

section = doc.sections[0]
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)


def no_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        borders.append(el)
    tblPr.append(borders)


def set_cell(cell, text, bold_lines=None, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    lines = text.split("\n")
    for idx, line in enumerate(lines):
        if idx > 0:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line)
        r.font.size = Pt(size)
        if bold_lines and idx in bold_lines:
            r.bold = True


def add_p(text="", space_after=6, indent=None, bold=False, tag=False, bullet=None):
    """Agrega un parrafo de texto normal (o una linea de tag jinja si tag=True)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0 if tag else space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if text:
        prefix = f"{bullet}  " if bullet else ""
        r = p.add_run(prefix + text)
        r.bold = bold
    return p


# Titulo
titulo = add_p(space_after=10)
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = titulo.add_run(
    "INFORME TECNICO N.\u00b0 {{ numero_informe }}-{{ anio }} "
    "{{ iniciales_tecnicos }}/{{ siglas_sede }}"
)
run.bold = True
run.underline = True
run.font.size = Pt(11)

# Tabla encabezado
tabla = doc.add_table(rows=5, cols=3)
tabla.autofit = False
widths = [Cm(2.5), Cm(0.5), Cm(11.5)]
for row in tabla.rows:
    for i, cell in enumerate(row.cells):
        cell.width = widths[i]

set_cell(tabla.rows[0].cells[0], "A")
set_cell(tabla.rows[0].cells[1], ":")
set_cell(tabla.rows[0].cells[2], "{{ destinatario_nombre }}\n{{ destinatario_cargo }}", bold_lines={1})

set_cell(tabla.rows[1].cells[0], "DE")
set_cell(tabla.rows[1].cells[1], ":")
set_cell(tabla.rows[1].cells[2], "{{ tecnico_nombre }}\n{{ tecnico_cargo }}", bold_lines={1})

set_cell(tabla.rows[2].cells[0], "T\u00c9CNICO")
set_cell(tabla.rows[2].cells[1], "")
set_cell(tabla.rows[2].cells[2], "")

set_cell(tabla.rows[3].cells[0], "ASUNTO")
set_cell(tabla.rows[3].cells[1], ":")
set_cell(tabla.rows[3].cells[2], "{{ asunto_linea1 }}\n{{ asunto_linea2 }}", bold_lines={0, 1})

set_cell(tabla.rows[4].cells[0], "FECHA")
set_cell(tabla.rows[4].cells[1], ":")
set_cell(tabla.rows[4].cells[2], "{{ fecha }}")

no_borders(tabla)

add_p(space_after=6)

linea = add_p(space_after=8)
pPr = linea._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'), 'single')
bottom.set(qn('w:sz'), '6')
bottom.set(qn('w:space'), '1')
bottom.set(qn('w:color'), 'auto')
pBdr.append(bottom)
pPr.append(pBdr)

add_p("Mediante el presente, se solicita {{ tipo_verbo }} {{ cantidad_equipos }} {{ tipo_equipo_plural }}.", space_after=8)

# --- BLOQUE REPETIBLE POR EQUIPO ---
add_p("{% for eq in equipos %}", tag=True)

add_p("{{ eq.texto_justificacion }}", space_after=8)

add_p(
    "Por tal motivo, se solicita la asignaci\u00f3n de un equipo nuevo, "
    "acorde a la planilla del personal indicado.",
    space_after=8,
)

add_p("RECOMENDACIONES:", space_after=4, bold=True)

add_p("{{ eq.cantidad }} {{ eq.tipo_bien }} con las siguientes caracter\u00edsticas:", space_after=2, indent=1, bullet="\u27a2")

add_p("{% for spec in eq.especificaciones %}", tag=True)
add_p("{{ spec }}", space_after=2, indent=2, bullet="\u2022")
add_p("{% endfor %}", tag=True)

add_p(space_after=6)

tabla2 = doc.add_table(rows=3, cols=3)
w2 = [Cm(2.3), Cm(0.5), Cm(11.7)]
for row in tabla2.rows:
    for i, cell in enumerate(row.cells):
        cell.width = w2[i]
set_cell(tabla2.rows[0].cells[0], "\u00c1rea")
set_cell(tabla2.rows[0].cells[1], ":")
set_cell(tabla2.rows[0].cells[2], "{{ eq.area }}")
set_cell(tabla2.rows[1].cells[0], "Usuario")
set_cell(tabla2.rows[1].cells[1], ":")
set_cell(tabla2.rows[1].cells[2], "{{ eq.usuario }}")
set_cell(tabla2.rows[2].cells[0], "Jefatura")
set_cell(tabla2.rows[2].cells[1], ":")
set_cell(tabla2.rows[2].cells[2], "{{ eq.jefatura }}")
no_borders(tabla2)

add_p("{% if not loop.last %}", tag=True)
pb = doc.add_paragraph()
brk = OxmlElement('w:br')
brk.set(qn('w:type'), 'page')
pb._p.append(brk)
add_p("{% endif %}", tag=True)
add_p("{% endfor %}", tag=True)
# --- FIN BLOQUE REPETIBLE ---

firma1 = add_p("_______________________", space_after=2)
firma1.alignment = WD_ALIGN_PARAGRAPH.CENTER

firma2 = add_p(space_after=2)
firma2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = firma2.add_run("{{ jefe_ti_nombre }}")
r.bold = True

firma3 = add_p()
firma3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = firma3.add_run("{{ jefe_ti_cargo }}")
r3.bold = True

RUTA_SALIDA.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(RUTA_SALIDA))
print(f"Plantilla creada en {RUTA_SALIDA}")
